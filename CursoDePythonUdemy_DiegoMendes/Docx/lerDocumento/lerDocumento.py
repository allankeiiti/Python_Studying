'''
Documentos no Word - .DOCX - Lendo o documento
Curso de Python

Python-docx
'''

import docx

doc = docx.Document('exemplo.docx')

#Lendo parágrafos do documento
print('Quantos parágrafos tem o documento: {}\n'.format(len(doc.paragraphs)))
for i in range(len(doc.paragraphs)):
    print(doc.paragraphs[i].text)

print('=====================================')
for i in range(len(doc.paragraphs)):
    print(len(doc.paragraphs[i].runs))
