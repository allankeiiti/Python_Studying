import docx

doc = docx.Document()
doc.add_paragraph('Curso de Python')
doc.add_paragraph('Escrevendo no documento por script python')
doc.paragraphs[0].style = 'Title'
doc.save('Doc1.docx')