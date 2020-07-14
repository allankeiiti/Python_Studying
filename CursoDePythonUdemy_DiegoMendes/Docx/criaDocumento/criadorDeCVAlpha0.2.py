import docx
from docx.shared import Pt
doc = docx.Document()

nome = input('Digite o nome a ser disponibilizado no currículo: ')

linha = doc.add_paragraph(nome)
linha.add_run('Curriculum Vitae')
linha.runs[0].underline = True
doc.save('Curriculo.docx')